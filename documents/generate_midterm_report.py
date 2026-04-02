"""Generate the midterm report as a Word document in ACL-style formatting."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# --- Page setup (ACL-style: letter, 1-inch margins) ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Automated Research Gap Discovery System\nMidterm Report')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Hirak Desai, Jaanaki Dave, Riya Verma, Shreeya Halwasia, Shubham Wani')
run.font.size = Pt(11)
run.font.name = 'Times New Roman'

affiliation = doc.add_paragraph()
affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affiliation.add_run('University of Southern California')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p


# ============================================================
# SECTION 1: REPORT ON TASKS PERFORMED
# ============================================================
add_heading('I. Report on Tasks Performed', level=1)

add_para(
    'Since the project proposal, we have implemented all five core pipeline stages described in our '
    'original design, moving from conceptual architecture to a fully functional end-to-end system. '
    'The following summarizes the concrete work completed.'
)

add_heading('Paper Ingestion', level=2)
add_para(
    'We implemented the ArxivClient module, which queries the ArXiv API with keyword-based search, '
    'downloads PDFs, and caches them locally to avoid redundant fetches. We additionally implemented '
    'the SemanticScholarClient, which enriches each paper with structured metadata (citation counts, '
    'venue, year) via the Semantic Scholar public API and discovers related papers through citation '
    'graph traversal. Top-cited discovered papers with ArXiv IDs are automatically downloaded, '
    'expanding the corpus beyond the initial query results. The system currently supports configurable '
    'corpus sizes of up to 100 papers per query.'
)

add_heading('Text Extraction and Segmentation', level=2)
add_para(
    'The PDFParser module extracts raw text from each PDF using PyMuPDF, preserving paragraph structure '
    'while normalizing whitespace. Text is truncated at the References section using word-boundary '
    'matching to avoid false splits. The SectionSegmenter identifies standard paper sections (Abstract, '
    'Introduction, Methods, Results, Discussion, Conclusion, Related Work, Background, Evaluation, and '
    'Experiments) using anchored regex patterns that match only at line starts with optional numbering, '
    'preventing false matches when section keywords appear in body text. SpaCy\'s en_core_web_sm model '
    'handles sentence tokenization.'
)

add_heading('Atomic Claim Extraction', level=2)
add_para(
    'We replaced the initial FLAN-T5-small prototype with Mistral 7B running locally via Ollama. '
    'The claim extraction prompt uses few-shot SciFact examples that demonstrate decomposing compound '
    'sentences into atomic, independently verifiable claims. Each extracted claim carries source '
    'provenance metadata (paper title, section name, original sentence) to enable citation tracing '
    'in downstream stages. A multi-stage filter removes metadata artifacts (DOIs, author names, URLs), '
    'short fragments, and non-claim text. In a test run on 3 papers, the system extracted 442 claims.'
)

add_heading('Embedding and Clustering', level=2)
add_para(
    'All claims are embedded into 768-dimensional vectors using SPECTER2 (Singh et al., 2023), a '
    'scientific document embedding model trained on 6M citation triplets across 23 fields. Claims '
    'are clustered using HDBSCAN with configurable parameters (min_cluster_size=2, '
    'cluster_selection_epsilon=0.8), which naturally handles variable-density regions without '
    'requiring a pre-specified cluster count. In our test run, 442 claims formed 102 clusters.'
)

add_heading('Contradiction Detection via NLI', level=2)
add_para(
    'Within each cluster, we apply pairwise NLI using a DeBERTa-v3-based cross-encoder '
    '(MoritzLaurer/deberta-v3-base-zeroshot-v1). A cosine similarity pre-filter (0.7\u20130.9 window) '
    'selects claim pairs that are related but not identical before running the more expensive NLI '
    'classification. The label order is read dynamically from the model\'s configuration rather than '
    'hardcoded, eliminating a critical bug identified during development where mismatched label ordering '
    'would have silently produced incorrect classifications.'
)

add_heading('Limitation Retrieval and Gap Generation', level=2)
add_para(
    'We implemented seed-query-based limitation retrieval as described in the proposal. Seed queries '
    '(e.g., "unresolved problem," "our approach does not handle," "limitation of this work") are '
    'embedded with SPECTER2 and compared against all claim embeddings via cosine similarity. Claims '
    'exceeding the similarity threshold are re-clustered to identify recurring limitation themes. '
    'The gap generation stage uses Mistral 7B to synthesize research gap statements from three '
    'sources: topic clusters (open questions), contradiction pairs, and limitation clusters. Each '
    'generated gap includes citation pointers to source papers and supporting claims, ensuring '
    'verifiability as specified in the proposal.'
)

add_heading('Evaluation Framework', level=2)
add_para(
    'We implemented the SciFact evaluation module with two automated metrics: (1) claim extraction '
    'quality measured by precision, recall, and F1 against SciFact ground-truth claims using BERTScore '
    'for semantic matching, and (2) NLI contradiction detection F1 on SciFact REFUTE pairs. The '
    'evaluation runs independently of the main pipeline and outputs results to JSON for analysis. '
    'The human evaluation framework (Likert scale for novelty, groundedness, and actionability) and '
    'baseline comparison against naive LLM prompting are designed and will be executed in the final phase.'
)

add_heading('User Interface', level=2)
add_para(
    'We built a Streamlit web application that provides a complete interactive frontend. The UI '
    'displays paper overviews (title, abstract preview, citation count, sections found) immediately '
    'after download, allowing users to browse papers while the pipeline runs. A sidebar shows live '
    'pipeline stage progress with checkmarks for completed stages. Results are organized into tabbed '
    'views (All Gaps, Open Questions, Contradictions, Limitations) with source papers and supporting '
    'claims for each gap. A cluster explorer allows inspection of the underlying claim groupings.'
)

# ============================================================
# SECTION 2: POTENTIAL RISKS AND CHALLENGES
# ============================================================
add_heading('II. Potential Risks and Challenges', level=1)

add_para(
    'Claim extraction quality. ',
    bold_prefix='1. '
)
doc.paragraphs[-1].runs[-1].text = (
    'Mistral 7B running locally on CPU produces acceptable but imperfect claims. Some extracted '
    'claims are overly verbose, repeat the input sentence, or include metadata artifacts that survive '
    'filtering. The quality varies across paper domains and writing styles. Since claim quality '
    'propagates through the entire pipeline (embedding, clustering, NLI, gap generation), extraction '
    'errors compound downstream.'
)

add_para(
    'Scalability of pairwise NLI. ',
    bold_prefix='2. '
)
doc.paragraphs[-1].runs[-1].text = (
    'Contradiction detection requires pairwise comparison within each cluster. While the cosine '
    'similarity pre-filter reduces the comparison space, large clusters can still produce a quadratic '
    'number of pairs. Processing 50+ papers generates hundreds of claims per cluster, making NLI '
    'inference the primary bottleneck. On a MacBook without GPU, each NLI call takes approximately '
    '0.5 seconds, and total NLI time scales with corpus size.'
)

add_para(
    'ArXiv API rate limiting. ',
    bold_prefix='3. '
)
doc.paragraphs[-1].runs[-1].text = (
    'ArXiv enforces rate limits that restrict the number of concurrent requests. When fetching 50\u2013100 '
    'papers, the system occasionally receives HTTP 429 responses. While we have implemented retry logic '
    'with exponential backoff, aggressive rate limiting during peak hours can still cause pipeline delays '
    'or failures on the first stage.'
)

add_para(
    'HDBSCAN sensitivity to embedding quality. ',
    bold_prefix='4. '
)
doc.paragraphs[-1].runs[-1].text = (
    'Clustering results are highly sensitive to the quality and distribution of claim embeddings. '
    'Noisy or overly generic claims can create large, incoherent clusters that dilute the specificity '
    'of downstream gap generation. The current HDBSCAN parameters (min_cluster_size=2, epsilon=0.8) '
    'were tuned for small corpora and may need adjustment for larger runs.'
)

add_para(
    'Gap generation hallucination. ',
    bold_prefix='5. '
)
doc.paragraphs[-1].runs[-1].text = (
    'Although gap generation is grounded in structurally discovered claims, the final LLM synthesis '
    'step can still produce gaps that are plausible-sounding but not well-supported by the input claims. '
    'The citation pointers help reviewers verify each gap, but automated detection of hallucinated '
    'gaps remains an open challenge.'
)

# ============================================================
# SECTION 3: RISK MITIGATION PLAN
# ============================================================
add_heading('III. Risk Mitigation Plan', level=1)

add_para(
    'For claim extraction quality, we plan to run the SciFact evaluation suite to quantify extraction '
    'precision and recall, then iteratively refine the few-shot prompt based on common error patterns. '
    'We will also experiment with increasing the temperature and sampling multiple extractions per '
    'sentence to improve coverage, followed by deduplication.'
)

add_para(
    'For NLI scalability, we will implement batch inference to process multiple claim pairs per '
    'forward pass, reducing per-pair overhead. We will also tighten the cosine similarity window '
    'to further prune candidate pairs before NLI, and investigate capping the maximum cluster size '
    'to bound the quadratic cost.'
)

add_para(
    'For ArXiv rate limiting, the retry-with-backoff mechanism is already in place. We will additionally '
    'implement local caching of API responses (not just PDFs) so that re-runs on the same query skip '
    'the API entirely. For large corpora, we will batch requests with configurable delays.'
)

add_para(
    'For clustering sensitivity, we plan to perform a parameter sweep over HDBSCAN settings '
    '(min_cluster_size, min_samples, epsilon) using the SciFact dataset as ground truth, selecting '
    'parameters that maximize cluster coherence as measured by intra-cluster cosine similarity.'
)

add_para(
    'For gap hallucination, we will implement a verification step that checks whether each generated '
    'gap is entailed by its supporting claims using the same NLI model. Gaps that are not supported '
    'will be flagged or regenerated. The baseline comparison (naive LLM prompting vs. structured '
    'pipeline) will quantify the hallucination reduction achieved by the claim-first approach.'
)

add_para(
    'The remaining work for project completion includes: (1) running the full SciFact evaluation and '
    'reporting quantitative results, (2) conducting the human evaluation study with Likert ratings, '
    '(3) executing the baseline comparison experiment, (4) tuning pipeline parameters on a validation '
    'set, and (5) writing the final report with experimental results and analysis.'
)

# ============================================================
# SECTION 4: INDIVIDUAL CONTRIBUTIONS
# ============================================================
add_heading('IV. Individual Contributions', level=1)

table = doc.add_table(rows=6, cols=2, style='Table Grid')
table.columns[0].width = Inches(1.5)
table.columns[1].width = Inches(5)

headers = table.rows[0].cells
headers[0].text = 'Member'
headers[1].text = 'Contributions'
for cell in headers:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

contributions = [
    ('Hirak Desai', 'Ollama/Mistral setup, SPECTER2 embeddings, configuration system, '
     'Streamlit UI with live progress tracking, ArXiv client with retry logic, end-to-end testing'),
    ('Jaanaki Dave', 'ArXiv API client implementation, PDF parsing and text extraction, '
     'section segmentation, initial pipeline prototyping'),
    ('Riya Verma', 'Claim extraction prompt engineering, SciFact few-shot example design, '
     'claim filtering heuristics, evaluation framework design'),
    ('Shreeya Halwasia', 'NLI contradiction detection module, DeBERTa-v3 integration, '
     'cosine similarity pre-filtering, label order validation'),
    ('Shubham Wani', 'Pipeline architecture and integration, HDBSCAN clustering module, '
     'Semantic Scholar client, limitation retrieval via seed queries, gap generation prompts'),
]

for i, (name, contrib) in enumerate(contributions):
    row = table.rows[i + 1].cells
    row[0].text = name
    row[1].text = contrib
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

# ============================================================
# REFERENCES
# ============================================================
add_heading('References', level=1)

refs = [
    '[1] Wadden, D., Lin, S., Lo, K., Wang, L. L., van Zuylen, M., Cohan, A., & Hajishirzi, H. '
    '(2020). Fact or fiction: Verifying scientific claims. In Proceedings of EMNLP 2020.',

    '[2] Wadden, D., Lo, K., Wang, L. L., Kuehl, S., Cohan, A., & Hajishirzi, H. (2022). '
    'SciFact-Open: Towards open-domain scientific claim verification. In Findings of EMNLP 2022.',

    '[3] Baek, J., Jauhar, S. K., Cucerzan, S., & Hwang, S. J. (2024). ResearchAgent: Iterative '
    'research idea generation over scientific literature with large language models. In Proceedings of ACL 2024.',

    '[4] Liu, R., & Shah, N. B. (2023). ReviewerGPT? An exploratory study on using large language '
    'models for paper reviewing. arXiv preprint arXiv:2306.00622.',

    '[5] Singh, A., et al. (2023). SciRepEval: A Multi-Format Benchmark for Scientific Document '
    'Representations. In Proceedings of ACL 2023.',

    '[6] Jiang, A. Q., et al. (2023). Mistral 7B. arXiv preprint arXiv:2310.06825.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/hirakdesai/Developer/USC/NLP-project/documents/Midterm_Report.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
